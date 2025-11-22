# cv_parser_fixed.py
import os
import json
import docx
import pdfplumber
from groq import Groq
import time
import warnings

# --- Settings ---
# GROQ_API_KEY should be passed from streamlit_app.py
INPUT_FOLDER = 'resumes'
OUTPUT_FOLDER = 'output'
FEEDBACK_FILE = 'feedback_examples.json'
CHUNK_SIZE_NORMAL = 4000  # Reduced significantly
CHUNK_SIZE_EMERGENCY = 2500  # Even smaller for emergency
MAX_RETRIES_CONNECTION = 3  # Reduced retries
RETRY_DELAY = 5  # Reduced delay

# --- Groq Client ---
try:
    client = Groq(api_key=GROQ_API_KEY)
    print("✅ Groq client initialized successfully.")
except Exception as e:
    print(f"❌ Error initializing Groq client: {e}")
    client = None

# --- Text Extraction ---
def extract_text(file_path, use_ocr=False):
    try:
        if file_path.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
            # Add table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += '\n' + cell.text
            return text
        elif file_path.lower().endswith('.pdf'):
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                with pdfplumber.open(file_path) as pdf:
                    text = ''
                    for page in pdf.pages:
                        if page:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + '\n'
                    return text
    except Exception as e:
        print(f"   ❌ Error extracting text from {os.path.basename(file_path)}: {e}")
    return ""

# --- Simplified Analysis Prompt ---
def build_simplified_prompt(text):
    return f"""
You are an expert HR analyst. Analyze this resume and provide a detailed JSON response.

RESUME TEXT:
{text}

Extract and analyze the following information:

1. PERSONAL INFO: Name, birth date, residence
2. EXPERIENCE: Total years, relevant years, types
3. EDUCATION: Engineering degrees and relevant education
4. CERTIFICATIONS: Professional certifications (CSWIP, ASNT, API, etc.)
5. INDUSTRY COMPETENCY: Match against these industries:
   - Oil & Gas Industry
   - Energy Industry  
   - Shipyard and Marine Industry
   - Wind and Turbine Industry
   - Power Generation Industry
   - Construction Industry
   - Manufacturing Industry
   - Automotive Industry
   - Aerospace Industry
   - Chemical Industry
   - Mining Industry
   - Nuclear Industry
   - Pharmaceutical Industry
   - Food and Beverage Industry
   - Textile Industry
   - Electronics Industry

6. EQUIPMENT COMPETENCY: Match against these equipment types:
   - Plate, Pipes and Fittings
   - Pressure Vessel
   - Storage Tanks
   - Heat Exchangers
   - Boilers
   - Turbines
   - Pumps and Compressors
   - Valves
   - Electrical Equipment
   - Instrumentation
   - Structural Steel
   - Cranes and Lifting Equipment
   - Conveyor Systems
   - HVAC Systems
   - Transformers
   - Generators
   - Motors
   - Control Systems
   - Safety Systems
   - Welding Equipment
   - Testing Equipment
   - Measurement Tools
   - Calibration Equipment
   - Fire Protection Systems
   - Environmental Systems
   - Process Equipment
   - Material Handling Equipment

SCORING SYSTEM:
- Experience: 30 points per year of relevant experience
- Education: 100 points for relevant engineering degree
- Certifications: CSWIP (30 pts), ASNT Level II (15 pts), ASNT Level III (60 pts), API (40 pts), IWE/IWI (60 pts), NACE (20 pts)

QUALIFICATION LEVELS:
- Highly Qualified: Score >= 350
- Qualified: 150-349
- Low Qualified/Junior: 50-149
- Rejected: < 50

Provide response in this exact JSON format:
{{
  "extracted_data": {{
    "name": "Full Name",
    "birth_date": "YYYY-MM-DD or null",
    "residence": "City, Country or null",
    "total_experience_years": 0,
    "relevant_experience_years": 0,
    "experience_type": ["Type1", "Type2"],
    "education": ["Degree - Major"],
    "qualifications": ["Cert1", "Cert2"],
    "competency_industry": ["Industry1", "Industry2"],
    "competency_equipment": ["Equipment1", "Equipment2"]
  }},
  "analysis": {{
    "calculated_score": 0,
    "qualification_level": "Level",
    "summary": "Brief professional summary",
    "justification": {{
      "qualification_level": "Explanation of level",
      "score_breakdown": "Experience: X, Education: Y, Certifications: Z, Total: Sum"
    }},
    "key_strengths": ["Strength1", "Strength2"],
    "areas_for_improvement": ["Area1", "Area2"]
  }}
}}

Be thorough but concise. Focus on relevant experience and certifications.
"""

# --- Fixed Chunking and Analysis ---
def chunk_text(text, chunk_size):
    """Split text into chunks of specified size."""
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i + chunk_size]
        chunks.append(chunk)
    return chunks

def analyze_with_groq(text, chunk_size):
    """Analyze text with Groq API."""
    if not client:
        return {"error": "Groq client not initialized"}
    
    # Split text into chunks if needed
    text_chunks = chunk_text(text, chunk_size)
    
    if len(text_chunks) > 1:
        print(f"   ℹ️ Splitting into {len(text_chunks)} chunks.")
    
    # Process each chunk
    for i, chunk in enumerate(text_chunks):
        print(f"   📝 Processing chunk {i+1}/{len(text_chunks)}")
        
        # Add delay between chunks
        if i > 0:
            print(f"   ⏳ Waiting 5 seconds before next chunk...")
            time.sleep(5)
        
        # Retry logic for each chunk
        for attempt in range(MAX_RETRIES_CONNECTION):
            try:
                prompt = build_simplified_prompt(chunk)
                
                # Make API call
                result = client.chat.completions.create(
                    messages=[
                        {
                            "role": "system", 
                            "content": "You are an expert HR analyst. Provide only valid JSON responses."
                        },
                        {
                            "role": "user", 
                            "content": prompt
                        }
                    ],
                    model="llama3-70b-8192",
                    response_format={"type": "json_object"},
                    temperature=0.1,
                    max_tokens=4000  # Reduced token limit
                )
                
                # Parse and return result
                response_json = json.loads(result.choices[0].message.content)
                return response_json
                
            except Exception as e:
                error_message = str(e).lower()
                
                # Check for rate limit or size errors
                if "rate_limit_exceeded" in error_message or "413" in error_message:
                    print(f"   ⚠️ Rate limit or size error: {e}")
                    return {"error": f"API Error: {e}", "is_size_error": True}
                
                # Retry for other errors
                if attempt < MAX_RETRIES_CONNECTION - 1:
                    delay = RETRY_DELAY * (attempt + 1)
                    print(f"   ⚠️ Connection error, retrying in {delay}s... (Attempt {attempt+1}/{MAX_RETRIES_CONNECTION})")
                    time.sleep(delay)
                else:
                    print(f"   ❌ All retries failed: {e}")
                    return {"error": f"Connection retries failed: {e}"}
    
    # If we get here, all chunks failed
    return {"error": "All chunks failed to process."}

def process_file(filename, chunk_size, use_ocr=False):
    """Process a single resume file."""
    file_path = os.path.join(INPUT_FOLDER, filename)
    
    # Extract text
    text = extract_text(file_path, use_ocr=use_ocr)
    if not text or not text.strip():
        print("   ⚠️ No text could be extracted.")
        return False, False
    
    print(f"   📄 Extracted {len(text)} characters from resume")
    print(f"   🔍 Analyzing with chunk size {chunk_size}...")
    
    # Analyze with Groq
    result = analyze_with_groq(text, chunk_size)
    
    # Check for errors
    if "error" in result:
        print(f"   ❌ AI analysis failed: {result['error']}")
        return False, result.get("is_size_error", False)
    
    # Validate result structure
    if not isinstance(result, dict) or 'analysis' not in result:
        print(f"   ❌ Invalid JSON structure.")
        return False, False
    
    # Add metadata
    result['metadata'] = {
        'source_file': filename,
        'processing_date': time.strftime("%Y-%m-%d %H:%M:%S"),
        'used_ocr': use_ocr,
        'chunk_size_used': chunk_size,
        'analysis_version': 'simplified_v1'
    }
    
    # Save result
    base_name = os.path.splitext(filename)[0]
    output_path = os.path.join(OUTPUT_FOLDER, f"{base_name}_analysis.json")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=4, ensure_ascii=False)
    
    print(f"   ✅ Analysis saved to: {output_path}")
    return True, False

def main():
    """Main function to process all resumes."""
    if not client:
        print("❌ Aborting due to Groq client initialization failure.")
        return
    
    # Create directories
    if not os.path.exists(INPUT_FOLDER):
        os.makedirs(INPUT_FOLDER)
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
    
    # Find resume files
    files = [f for f in os.listdir(INPUT_FOLDER) if f.lower().endswith(('.pdf', '.docx'))]
    
    if not files:
        print("❌ No resume files found in the input folder.")
        return
    
    print(f"🚀 Starting Resume Analysis: Found {len(files)} file(s)")
    
    failed_files = []
    
    # Process each file
    for index, filename in enumerate(files):
        print(f"\n[{index + 1}/{len(files)}] Processing: {filename}")
        
        # Try with normal chunk size
        success, is_size_error = process_file(filename, CHUNK_SIZE_NORMAL)
        
        # If failed due to size, try with emergency chunk size
        if not success and is_size_error:
            print("   ⚠️ Retrying with smaller chunk size...")
            success, is_size_error = process_file(filename, CHUNK_SIZE_EMERGENCY)
        
        # If still failed, add to failed files
        if not success:
            failed_files.append(filename)
        
        # Add delay between files
        if index < len(files) - 1:
            print("   ⏳ Waiting 3 seconds before next file...")
            time.sleep(3)
    
    # Print summary
    print("\n" + "="*50)
    print("📊 PROCESSING SUMMARY")
    print("="*50)
    
    if failed_files:
        print(f"❌ Failed to process {len(failed_files)} file(s):")
        for filename in failed_files:
            print(f"   - {filename}")
    else:
        print("🎉 All files processed successfully!")
    
    successful_files = len(files) - len(failed_files)
    print(f"✅ Successfully processed: {successful_files}/{len(files)} files")

if __name__ == "__main__":
    main()